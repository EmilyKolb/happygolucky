from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re
import tkinter as tk
from tkinter import filedialog
from tkinter import simpledialog

#to do: err toward keeping in slate
#to do: handling subtitles?

def changewordstodigits(text):

    dictionary = {}
    zerototeensdictionary = {

        'zero':'0',
        'one':'1',
        'two':'2',
        'three':'3',
        'four':'4',
        'five':'5',
        'six':'6',
        'seven':'7',
        'eight':'8',
        'nine':'9',
        'ten':'10',

        'eleven':'11',
        'twelve':'12',
        'thirteen':'13',
        'fourteen':'14',
        'fifteen':'15',
        'sixteen':'16',
        'seventeen':'17',
        'eighteen':'18',
        'nineteen':'19',
        'twenty':'20',
        'thirty':'30',
        'forty':'40',
        'fifty':'50',
        'sixty':'60',
        'seventy':'70',
        'eighty':'80',
        'ninety':'90'
        }


    for character in text:
        zerodictionarykeys = list(zerototeensdictionary.keys())
        zerodictionaryvalues = list(zerototeensdictionary.values())
        punctlist = [',','.','?','-']
        punctcounter = 0
        dictcounter = 0

        while punctcounter < len(punctlist):
            punct = punctlist[punctcounter]
            dictcounter = 0
            while dictcounter < (len(zerodictionarykeys)):
                alphaword = zerodictionarykeys[dictcounter]
                numaword = zerodictionaryvalues[dictcounter]
                dictcounter += 1
                text = text.replace(" " + alphaword + " "," " + numaword + " ")
                text = text.replace(alphaword + punct,numaword + punct)

            punctcounter += 1


        return(text)


def teenstodigits(text):
    counter = 0
    for character in text:
        counter += 1


        if counter + 3 < len(text) and text[counter].isdigit() and text[counter + 1] == '0' and text[counter + 2] == ' ' and text[counter + 3].isdigit():
            start = counter + 1
            end = counter + 3

            text = text[:start] + text[end:]

    return(text)

def getridofhundreds(text):
    counter = 0

    for character in text:
        counter += 1

        foundhundredand = text.find('hundred and')
        foundhundred = text.find('hundred ')
        foundhundreds = text.find('hundreds')
        foundoh = text.find('oh')


        if counter + 13 < len(text) and text[counter].isdigit() and text[counter + 1] == ' ' and (counter + 2) == foundhundredand and text[counter + 14].isdigit():
            start = counter + 1
            end = foundhundredand + 12

            text = text[:start] + "0" + text[end:]

        elif counter + 3 < len(text) and text[counter].isdigit() and text[counter + 1] == ' ' and (counter + 2) == foundhundred:

            start = counter + 1
            end = foundhundred + 8

            text = text[:start] + "0" + text[end:]

        elif counter + 4 < len(text) and text[counter].isdigit() and (counter + 2) == foundhundreds: #turns 18 hundreds (eighteen hundreds before def(changewordstodigits)) into 1800s

            start = counter + 1
            end = counter + 10
            text = text[:start] + "00s" + text[end:]

        elif counter +2 < len(text) and text[counter].isdigit() and (text[counter + 1] == ' ' or text[counter + 1] == '-') and counter + 2 == foundoh:
            text = text[:counter + 1] + "0" + text[counter + 5:]

    return(text)

def combinenumbers(text):
    counter = 0

    for character in text:
         counter += 1
         if counter + 2 < len(text) and text[counter].isdigit() and text[counter + 1] == ' ' and text[counter + 2].isdigit():
             text = text[:counter + 1] + text[counter + 2:]

    return(text)



def remove_bracketed_vo_notes(text):

    while '\n\n\n' in text: #remove consecutive newlines beyond 2
        text = text.replace('\n\n\n','\n\n')

    paraList = text.split('\n\n') #split into paragraphs

    for i in range(len(paraList)): #remove comments
        para = paraList[i]

        while '//' in para: #remove '//'
            start = para.find('//')
            nextUpper = 9999

            counter = 0
            for character in para[(start + 2):]:
                counter += 1
                if (start + 4 + counter) < len(para) and character == '/' and para[start + 2 + counter] == '/' and para[start + 3 + counter] == '\n' and para[start + 4 + counter].isupper() and nextUpper == 9999:
                    nextUpper = (start + 4 + counter)
            if para[(start + 2):].find('//') == -1: #if multiple '//' in paragraph
                end = min(len(para), nextUpper)
            else:
                end = min(len(para), para[(start + 2):].find('//') + (start + 4), nextUpper)

            para = para[:start] + para[end:]

        while '[' in para: #remove '[...]'
            start = para.find('[')
            if para.find(']') == -1:
                end = len(para)
            else:
                end = para.find(']')
            para = para[:start] + para[end + 1:]

        while '(' in para: #remove '(...)'
            start = para.find('(')
            if para.find(')') == -1:
                end = len(para)
            else:
                end = para.find(')')
            para = para[:start] + para[end + 1:]

        counter = 0
        afterTimeCode = False
        for character in para[:-1]: #eliminate unwanted newlines -- !!! not quite perfect !!!
            counter += 1
            if (character == '.' or character == '!' or character == '?' or character == ':' or character == ' ') and para[counter] == '\n' and afterTimeCode:
                para = para
            elif para[counter] == '\n' and afterTimeCode:
                para = para[:counter] + ' ' + para[(counter + 1):]
            elif character == '\n' and not afterTimeCode and para[counter - 2].isdigit():
                para = para[:counter] + '\n' + para[counter:]
                afterTimeCode = True
            elif (not (character == '.' or character == '!' or character == '?' or character == ':' or character == ' ')) and para[counter] == '\n' and not afterTimeCode:
                para = para[:counter] + ' ' + para[(counter + 1):]

        paraList[i] = para

    newText = ''

    for para in paraList: #recombine paragraphs
        newText += (para + '\n\n')

    return newText


def insert_hard_returns(text):
    counter = 0
    ret = ''
    length = len(text) - 1
    for character in text:
        counter += 1

        if character.isdigit() and text[counter + 2].isdigit() and text[counter - 2] == '\n':
            ret += '\n' + character

        elif counter < length and character == ' ' and text[counter] == '\n':
            ret = ret

        else:
            ret += character

    return ret

def delete_many_hard_returns(text):
    counter = 0
    newstring = ''
    ret = ''
    for character in text:
        counter += 1

        if character is "\n" and text[counter - 2] == "\n":
            ret = ret[:counter - 2]

        else:
            ret += character

    return ret


def delete_slate(text):
    ret = text
    counter = 0
    start = 0

    for character in text:
        counter += 1

        if character is "\n" and text[counter - 2] == "\n" and ret == text:
            ret = text[counter:]

    return ret

def modify_timecode(text):

    ret = ''

    for character in text:
        if character == "\t":
            ret += " - "
        else:
            ret += character

    return ret

def singledigitstowords(text,inverted_dict):
    zerotenkeys = list(inverted_dict.keys())
    zerotenvalues = list(inverted_dict.values())
    punctlist = [' ', '.','!','?']
    punctcounter = 0


    while punctcounter < len(punctlist):
        punct = punctlist[punctcounter]
        dictcounter = 0

        while dictcounter < (len(zerotenkeys)):
            numword = zerotenkeys[dictcounter]
            alphaword = zerotenvalues[dictcounter]
            dictcounter += 1
            text = text.replace(' ' + numword + punct,' ' + alphaword + punct)

        punctcounter += 1


    return(text)

def numsinwords(text,inverted_dict):
    counter = 0
    dictcounter = 0
    zerotenkeys = list(inverted_dict.keys())
    zerotenvalues = list(inverted_dict.values())
    newText = ""

    for character in text:
        counter += 1
        dictcounter = 0
        while dictcounter < (len(zerotenkeys)):
            numword = zerotenkeys[dictcounter]
            alphaword = zerotenvalues[dictcounter]
            dictcounter += 1

            if character == numword and text[counter - 2].isalpha():
                start = (counter - 1)
                end = counter

                newText = text[:start] + alphaword + text[end:]

    return(newText)

def getridofdashes(text):
    counter = 0
    temptext = ""

    for character in text:
        counter += 1
        if counter + 2 < len(text) and character == "-" and text[counter].isupper() and text[counter - 2].isupper():
            text = text
        else:
            temptext += character
    return(temptext)


def modify_to_docx(text, title):
    document = Document('C:/Scrypti/Program/ScryptiTemplate/ScryptiTemplate.docx') #needs a template document in order to work. Template document cannot be empty.
    paragraph = document.paragraphs[0]
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

    style = document.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    titleparagraph = document.add_paragraph()
    titleparagraph.add_run('Captionmax Video Description Script' + "\n").bold = True
    titleparagraph.add_run(title).bold = True

    titleparagraph_format = titleparagraph.paragraph_format
    titleparagraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



    descriptiontext = ''
    indescription = True
    counter = 0
    length = len(text)
    for character in text:
        counter += 1

        if character.isdigit() and text[counter].isdigit and text[counter + 1] == ':':
            indescription = True

        elif (counter < length and character == '\n' and text[counter] == '\n') or counter == length:
            indescription = False

        if not indescription:
            indescription = True
            descriptiontext = descriptiontext[1:]
            descriptionparagraph = document.add_paragraph(descriptiontext)
            descriptiontext = ''
            descriptionparagraph_format = descriptionparagraph.paragraph_format
            descriptionparagraph_format.keep_together = True

        else:
            descriptiontext += character


        filename = ''
        firstnumber = True
        for character in title:
            if character.isalpha():
                filename += character
            elif character.isdigit() and firstnumber:
                filename += '-' + character
                firstnumber = False
            elif character.isdigit():
                filename += character
            elif character == ' ':
                filename = filename
            elif character == '&':
                filename += '_and_'

    savetitle = filename + '-VDS.docx'

    document.save("C:/Scrypti/Completed Files/%s" % savetitle)



def main():

    zerotendict = {
    'zero':'0',
    'one':'1',
    'two':'2',
    'three':'3',
    'four':'4',
    'five':'5',
    'six':'6',
    'seven':'7',
    'eight':'8',
    'nine':'9',
    'ten':'10',
    }

    inverted_dict = dict([[v,k] for k,v in zerotendict.items()])


    root = tk.Tk()
    root.withdraw()
    file = filedialog.askopenfilename()

    root.withdraw()
    answer = simpledialog.askstring("File Name", "Episode Title (e.g. Chicago Med 315)",
                                    parent=root)
    root.withdraw()
    title = answer

    infile = open(file,"r")
    text = infile.read()

    text = "\n\n" + delete_slate(text)

    #these four functions handle turning "eighteen eighty four" into "1884" according to Chicago Manual of Style guidelines
    text = changewordstodigits(text)
    text = teenstodigits(text)
    text = getridofhundreds(text)
    text = combinenumbers(text)

    #handles deleting vo notes and formatting timecode
    text = remove_bracketed_vo_notes(text)
    text = delete_many_hard_returns(text)
    text = modify_timecode(text)
    text = insert_hard_returns(text)

    #more number handling
    text = singledigitstowords(text,inverted_dict)
    text = numsinwords(text,inverted_dict)

    #turns F-B-I to FBI
    text = getridofdashes(text)


    modify_to_docx(text, title)




if __name__ == '__main__':
    main()
